import io
from docx import Document
from docx.shared import Inches

def img_to_bytes(img, img_format="JPEG"):
    """
    Converts a PIL image to bytes format.

    Args:
        img (PIL.Image.Image): The image to be converted.
        img_format (str): The format of the image (e.g., JPEG, PNG).

    Returns:
        bytes: The byte representation of the image.
    """
    try:
        img_byte_arr = io.BytesIO()
        img.save(img_byte_arr, format=img_format)
        img_byte_arr.seek(0)  # Rewind to the beginning of the byte stream
        return img_byte_arr.read()  # Return the byte data
    except Exception as e:
        print(f"Error in converting image to bytes: {e}")
        return None

def convert_to_docx(content, images):
    """
    Converts content and images into a DOCX file.

    Args:
        content (str): The text content to include in the DOCX.
        images (list): List of image byte arrays to include in the DOCX.

    Returns:
        BytesIO: A byte stream containing the DOCX file.
    """


    try:
        doc = Document()
        doc.add_paragraph(content)

        for i, img_bytes in enumerate(images):
            img_stream = io.BytesIO(img_bytes)
            doc.add_picture(img_stream, width=Inches(3))

        doc_stream = io.BytesIO()
        doc.save(doc_stream)
        doc_stream.seek(0)
        return doc_stream
    except Exception as e:
        print(f"Error creating DOCX: {e}")
        return None
